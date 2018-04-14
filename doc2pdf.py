#-------------------------------------------------------------------------------
# Name:        module1
# Purpose:
#
# Author:      acer
#
# Created:     05-03-2018
# Copyright:   (c) acer 2018
# Licence:     <your licence>
#-------------------------------------------------------------------------------
import docx
doc = docx.Document('test.docx')
n = len(doc.paragraphs)
for i in range(0,n):
    print(doc.paragraphs[i].text)
